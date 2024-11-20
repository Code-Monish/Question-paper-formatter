from docx import Document
import sqlite3
import os
import re

db_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '../db/COURSES.db'))
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

def QuestionDOCX():
    cursor.execute("SELECT Question, Marks, BTL, CO FROM Questions")
    questions = cursor.fetchall()

    # Create a new Document
    doc = Document()
    doc.add_heading('Questions', level=1)
    # doc.add_section()
    for question, marks, btl, co in questions:
        # Add question
        p = doc.add_paragraph()
        p.add_run(question).bold = True
        # Add marks, BTL, and CO to the right of the question
        p.add_run(f' (Marks: {marks}, BTL: {btl}, CO: {co})')

    try:
        cursor.execute("SELECT file FROM Questions")
        file_name = cursor.fetchone()[0]  # Assuming first column of the result

        # Sanitize the filename
        safe_filename = re.sub(r'[<>:"/\\|?*]', '', file_name)
        safe_filename = safe_filename.replace('-', '_').replace(' ', '_')

        # Ensure the directory exists
        base_path = "../Back_End_Configuring/QuestionBank"
        os.makedirs(base_path, exist_ok=True)

        # Create full file path
        full_file_path = os.path.normpath(os.path.join(base_path, safe_filename + ".docx"))
        print(full_file_path)
        # Save the document
        doc.save(full_file_path)

        print(f"Document created successfully at: {full_file_path}")
        conn.close()

    except Exception as e:
        print(f"Error creating document: {e}")
        # Optional: log the error or handle it appropriately
        if conn:
            conn.close()
